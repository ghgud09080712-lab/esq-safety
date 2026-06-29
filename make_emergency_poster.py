from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"C:\Users\zxcas\Downloads\중대산업재해_비상대응_행동요령_포스터.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9D9D9", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)


def add_run(paragraph, text, size=12, bold=False, color="000000"):
    run = paragraph.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(0.8)
section.left_margin = Cm(0.8)
section.right_margin = Cm(0.8)

styles = doc.styles
styles["Normal"].font.name = "Malgun Gothic"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
add_run(title, "중대산업재해 비상대응 행동요령", 22, True, "C00000")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(8)
add_run(subtitle, "발생 시 또는 급박한 위험이 있을 때 즉시 실시", 12, True, "1F4E79")

alert = doc.add_table(rows=1, cols=1)
alert.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = alert.cell(0, 0)
set_cell_shading(cell, "FFF2CC")
set_cell_border(cell, "BF9000", "12")
set_cell_margins(cell, 160, 180, 160, 180)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "위험하다고 판단되면 누구든지 작업을 멈출 수 있습니다.", 13, True, "7F6000")

doc.add_paragraph().paragraph_format.space_after = Pt(2)

steps = [
    ("1", "작업중지", "즉시 작업을 멈추고 설비·장비를 안전상태로 전환"),
    ("2", "근로자 대피", "지정 대피로를 따라 집결장소로 이동"),
    ("3", "신고·보고", "관리감독자·안전부서·119에 즉시 연락"),
    ("4", "구조조치", "부상자 응급조치 및 구조 요청"),
    ("5", "위험요인 제거", "전기·가스·화학물질·화재 등 추가 위험 차단"),
    ("6", "추가 피해 방지", "출입 통제, 주변 작업 중지, 2차 사고 예방"),
]

table = doc.add_table(rows=len(steps), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_fixed(table, [850, 2200, 7600])
for row, (num, action, detail) in zip(table.rows, steps):
    n, a, d = row.cells
    for c in row.cells:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(c, "BFBFBF", "8")
        set_cell_margins(c, 95, 110, 95, 110)
    set_cell_shading(n, "C00000")
    set_cell_shading(a, "EAF2F8")
    clear_cell(n)
    clear_cell(a)
    clear_cell(d)
    pn = n.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(pn, num, 17, True, "FFFFFF")
    pa = a.paragraphs[0]
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(pa, action, 13, True, "1F4E79")
    pd = d.paragraphs[0]
    add_run(pd, detail, 10.5, False, "000000")

doc.add_paragraph().paragraph_format.space_after = Pt(2)

contact = doc.add_table(rows=4, cols=2)
contact.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_fixed(contact, [2050, 8600])
labels = [
    ("현장 책임자", "성명/연락처:"),
    ("안전관리자", "성명/연락처:"),
    ("비상연락", "119 / 사내 비상연락망:"),
    ("집결장소", "지정 집결지:"),
]
for i, (label, value) in enumerate(labels):
    left, right = contact.rows[i].cells
    set_cell_shading(left, "D9EAF7")
    for c in (left, right):
        set_cell_border(c, "BFBFBF", "8")
        set_cell_margins(c, 75, 110, 75, 110)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    clear_cell(left)
    clear_cell(right)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(lp, label, 9.5, True, "1F4E79")
    add_run(right.paragraphs[0], value, 9.5, False, "000000")

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.paragraph_format.space_before = Pt(8)
add_run(note, "게시 위치: 현장 게시판 · 작업허가서 발행 장소 · 위험작업 구역 인근", 9, True, "404040")

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.paragraph_format.space_before = Pt(2)
add_run(foot, "관련 절차: 비상대응 매뉴얼 및 비상조치계획서 / 교육·TBM 전파 기록 보관", 8.5, False, "666666")

doc.save(OUT)
print(OUT)
